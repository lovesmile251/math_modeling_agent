from __future__ import annotations

import hashlib
import json
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from tools.report_builder import (
    Block,
    CodeBlock,
    Document,
    Heading,
    ImageBlock,
    ListBlock,
    MathBlock,
    Paragraph,
    TableBlock,
    iter_inline_runs,
)
from tools.math_render import render_latex

# Default CJK-capable font names per backend. reportlab ships the Adobe
# STSong-Light CID font so Chinese renders without any external font file.
_DOCX_CJK_FONT = "SimSun"
_PDF_CJK_FONT = "STSong-Light"
_DOCX_CODE_FONT = "Consolas"
_MAX_TABLE_COLUMNS = 8
_LONG_TABLE_ROWS = 24
_LONG_CODE_LINE = 96
DEFAULT_DOCX_TEMPLATE_PATH = (
    Path(__file__).resolve().parent / "paper_templates" / "assets" / "national_contest_2025.docx"
)
DEFAULT_DOCX_TEMPLATE_SHA256 = "aedc429e0951fbc5d4ce92c0c18a6c28218782baaffe23307a97540388135bc5"


@dataclass(frozen=True)
class DocxTemplateField:
    role: str
    source: str
    docx_style: str
    template_anchor: str
    formatting: str


@dataclass(frozen=True)
class DocxTemplateLayoutCheck:
    passed: bool
    document_path: str
    template_path: str
    template_sha256: str
    field_mapping: list[dict[str, str]]
    metrics: dict[str, Any]
    warnings: list[str]

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


_DOCX_TEMPLATE_FIELD_MAPPING: tuple[DocxTemplateField, ...] = (
    DocxTemplateField(
        role="title",
        source="Document.title",
        docx_style="Normal",
        template_anchor="基于XXX模型的XXX问题研究",
        formatting="centered, 16pt, template section geometry",
    ),
    DocxTemplateField(
        role="abstract_heading",
        source="Heading text matching 摘要/Abstract",
        docx_style="Normal",
        template_anchor="摘要",
        formatting="centered, 14pt",
    ),
    DocxTemplateField(
        role="keywords",
        source="Heading/paragraph text matching 关键词/Keywords",
        docx_style="Normal",
        template_anchor="关键词：XXX，XXX，XXX，XXX",
        formatting="left aligned, 12pt",
    ),
    DocxTemplateField(
        role="section_heading",
        source="Heading(level=1)",
        docx_style="Heading 1",
        template_anchor="问题重述 / 问题分析 / 模型假设",
        formatting="template heading style",
    ),
    DocxTemplateField(
        role="subsection_heading",
        source="Heading(level=2)",
        docx_style="Heading 2",
        template_anchor="问题一的分析",
        formatting="template heading style",
    ),
    DocxTemplateField(
        role="body",
        source="Paragraph/ListBlock/CodeBlock text",
        docx_style="Normal/List Paragraph",
        template_anchor="正文示例段落",
        formatting="template fonts with explicit CJK fallback where needed",
    ),
    DocxTemplateField(
        role="display_math",
        source="MathBlock and inline math runs",
        docx_style="Normal",
        template_anchor="这里插入公式",
        formatting="rendered equation image, centered for display math",
    ),
    DocxTemplateField(
        role="table",
        source="TableBlock",
        docx_style="Table Grid",
        template_anchor="表1 本文的符号说明",
        formatting="split when wider than 8 columns; explicit cell fonts",
    ),
    DocxTemplateField(
        role="figure",
        source="ImageBlock",
        docx_style="Normal",
        template_anchor="图1 问题的总分析",
        formatting="max width 5.8in with centered caption paragraph",
    ),
)


def docx_template_field_mapping() -> list[dict[str, str]]:
    """Return the formal national-contest DOCX template field mapping."""

    return [asdict(item) for item in _DOCX_TEMPLATE_FIELD_MAPPING]


def _equation_dir(output_path: Path) -> Path:
    return output_path.parent / "_equations"


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def export_docx(
    doc: Document,
    output_path: Path,
    template_path: Path | None = DEFAULT_DOCX_TEMPLATE_PATH,
) -> Path:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    docx, using_template = _new_docx_document(DocxDocument, template_path)

    if not using_template:
        normal = docx.styles["Normal"]
        normal.font.name = _DOCX_CJK_FONT
        normal.font.size = Pt(11)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), _DOCX_CJK_FONT)

    eq_dir = _equation_dir(output_path)

    _docx_add_title(docx, doc.title, Pt, WD_ALIGN_PARAGRAPH, using_template)

    for block in _content_blocks(doc):
        if isinstance(block, Heading):
            heading = _docx_add_heading(docx, block, Pt, WD_ALIGN_PARAGRAPH, using_template)
            _docx_add_runs(heading, block.text, eq_dir)
            _docx_apply_heading_run_overrides(heading, block.text, Pt, using_template)
        elif isinstance(block, Paragraph):
            _docx_add_runs(docx.add_paragraph(), block.text, eq_dir)
        elif isinstance(block, MathBlock):
            # Render display math as image (works in all viewers, no OMML leak)
            image = render_latex(block.latex, eq_dir, display=True)
            if image is not None and image.path.exists():
                para = docx.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    para.add_run().add_picture(
                        str(image.path),
                        width=Inches(min(image.width_pt / 72, 5.8)),
                    )
                except Exception:
                    para.add_run(_math_fallback_text(block.latex))
            else:
                para = docx.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run(_math_fallback_text(block.latex))
        elif isinstance(block, ListBlock):
            for index, item in enumerate(block.items, start=1):
                paragraph = _docx_add_list_paragraph(docx, block.ordered, index)
                _docx_add_runs(paragraph, item, eq_dir)
        elif isinstance(block, CodeBlock):
            para = docx.add_paragraph()
            run = para.add_run(block.text)
            _docx_set_run_fonts(run, _DOCX_CODE_FONT, _DOCX_CJK_FONT)
            run.font.size = Pt(9)
        elif isinstance(block, ImageBlock):
            if block.path.exists():
                try:
                    docx.add_picture(str(block.path), width=Inches(5.8))
                except Exception:
                    docx.add_paragraph(f"[图片无法嵌入：{block.path.name}]")
                if block.caption:
                    cap = docx.add_paragraph(block.caption)
                    cap.alignment = 1  # centered
        elif isinstance(block, TableBlock):
            for caption, headers, rows in _split_table_columns(block):
                if caption:
                    docx.add_paragraph(caption)
                table = docx.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for idx, header in enumerate(headers):
                    table.rows[0].cells[idx].text = ""
                    _docx_add_runs(table.rows[0].cells[idx].paragraphs[0], header, eq_dir)
                for row in rows:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row[: len(headers)]):
                        cells[idx].text = ""
                        _docx_add_runs(cells[idx].paragraphs[0], value, eq_dir)
                _docx_format_table(table, len(headers), Inches, Pt)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    docx.save(str(output_path))
    return output_path


def _new_docx_document(docx_document_factory, template_path: Path | None):
    if template_path is not None:
        path = Path(template_path)
        if path.exists():
            docx = docx_document_factory(str(path))
            _docx_clear_body(docx)
            return docx, True
    return docx_document_factory(), False


def _docx_clear_body(docx) -> None:
    from docx.oxml.ns import qn

    body = docx._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def check_docx_template_layout(
    docx_path: Path,
    template_path: Path | None = DEFAULT_DOCX_TEMPLATE_PATH,
) -> DocxTemplateLayoutCheck:
    """Audit that a generated DOCX still follows the formal contest template."""

    warnings: list[str] = []
    metrics: dict[str, Any] = {}
    docx_path = Path(docx_path)
    template = Path(template_path) if template_path is not None else None

    if not docx_path.exists():
        return DocxTemplateLayoutCheck(
            passed=False,
            document_path=str(docx_path),
            template_path=str(template or ""),
            template_sha256="",
            field_mapping=docx_template_field_mapping(),
            metrics={},
            warnings=[f"DOCX file does not exist: {docx_path}"],
        )
    if template is None or not template.exists():
        return DocxTemplateLayoutCheck(
            passed=False,
            document_path=str(docx_path),
            template_path=str(template or ""),
            template_sha256="",
            field_mapping=docx_template_field_mapping(),
            metrics={},
            warnings=[f"DOCX template asset does not exist: {template}"],
        )

    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    template_hash = _file_sha256(template)
    metrics["template_hash_matches"] = int(template_hash == DEFAULT_DOCX_TEMPLATE_SHA256)
    if template_hash != DEFAULT_DOCX_TEMPLATE_SHA256:
        warnings.append("DOCX template hash does not match the registered 2025 national-contest asset")

    try:
        generated = DocxDocument(str(docx_path))
        template_doc = DocxDocument(str(template))
    except Exception as exc:
        return DocxTemplateLayoutCheck(
            passed=False,
            document_path=str(docx_path),
            template_path=str(template),
            template_sha256=template_hash,
            field_mapping=docx_template_field_mapping(),
            metrics=metrics,
            warnings=[f"DOCX layout audit failed to open document: {exc}"],
        )

    metrics["paragraphs"] = len(generated.paragraphs)
    metrics["tables"] = len(generated.tables)
    metrics["sections"] = len(generated.sections)
    metrics["field_mapping_count"] = len(_DOCX_TEMPLATE_FIELD_MAPPING)

    _audit_docx_section_geometry(generated, template_doc, metrics, warnings)
    _audit_docx_placeholders(generated, metrics, warnings)
    _audit_docx_title(generated, WD_ALIGN_PARAGRAPH, warnings)
    _audit_docx_tables(generated, metrics, warnings)

    return DocxTemplateLayoutCheck(
        passed=not warnings,
        document_path=str(docx_path),
        template_path=str(template),
        template_sha256=template_hash,
        field_mapping=docx_template_field_mapping(),
        metrics=metrics,
        warnings=warnings,
    )


def write_docx_template_layout_report(
    report: DocxTemplateLayoutCheck,
    output_path: Path,
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(report.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
    return output_path


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _audit_docx_section_geometry(generated, template_doc, metrics: dict[str, Any], warnings: list[str]) -> None:
    if not generated.sections:
        warnings.append("DOCX has no sections")
        return
    if not template_doc.sections:
        warnings.append("DOCX template has no sections")
        return
    generated_section = generated.sections[0]
    template_section = template_doc.sections[0]
    attrs = (
        "page_width",
        "page_height",
        "left_margin",
        "right_margin",
        "top_margin",
        "bottom_margin",
    )
    mismatches: list[str] = []
    for attr in attrs:
        actual = int(getattr(generated_section, attr))
        expected = int(getattr(template_section, attr))
        metrics[f"section_{attr}"] = actual
        metrics[f"template_{attr}"] = expected
        if actual != expected:
            mismatches.append(attr)
    if mismatches:
        warnings.append("DOCX section geometry differs from template: " + ", ".join(mismatches))


def _audit_docx_placeholders(generated, metrics: dict[str, Any], warnings: list[str]) -> None:
    text = _docx_visible_text(generated)
    placeholders = (
        "基于XXX模型的XXX问题研究",
        "关键词：XXX，XXX，XXX，XXX",
        "这里插入公式，务必用公式编辑器，不要截图！",
        "这里插入公式",
    )
    remaining = [item for item in placeholders if item in text]
    metrics["template_placeholder_hits"] = len(remaining)
    if remaining:
        warnings.append("DOCX still contains template placeholder text: " + ", ".join(remaining[:4]))


def _audit_docx_title(generated, WD_ALIGN_PARAGRAPH, warnings: list[str]) -> None:
    title = next((paragraph for paragraph in generated.paragraphs if paragraph.text.strip()), None)
    if title is None:
        warnings.append("DOCX has no visible title paragraph")
        return
    if title.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        warnings.append("DOCX title paragraph is not centered like the contest template")
    sizes = [run.font.size.pt for run in title.runs if run.font.size is not None]
    if sizes and not any(abs(size - 16.0) < 0.1 for size in sizes):
        warnings.append("DOCX title paragraph does not preserve the template 16pt title size")


def _audit_docx_tables(generated, metrics: dict[str, Any], warnings: list[str]) -> None:
    max_columns = 0
    for table in generated.tables:
        max_columns = max(max_columns, len(table.columns))
        if len(table.columns) > _MAX_TABLE_COLUMNS:
            warnings.append(f"DOCX table has {len(table.columns)} columns after export; expected split layout")
    metrics["max_table_columns"] = max_columns


def _docx_visible_text(docx) -> str:
    parts = [paragraph.text for paragraph in docx.paragraphs]
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _docx_add_title(docx, title: str, Pt, WD_ALIGN_PARAGRAPH, using_template: bool):
    if not using_template:
        return docx.add_heading(title, level=0)

    paragraph = docx.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(7.85)
    paragraph.paragraph_format.space_after = Pt(7.85)
    run = paragraph.add_run(title)
    run.font.size = Pt(16)
    return paragraph


def _docx_add_heading(docx, block: Heading, Pt, WD_ALIGN_PARAGRAPH, using_template: bool):
    text = block.text.strip()
    if using_template and text == "摘要":
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(7.85)
        paragraph.paragraph_format.space_after = Pt(7.85)
        return paragraph
    if using_template and text == "关键词":
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(7.85)
        paragraph.paragraph_format.space_after = Pt(7.85)
        return paragraph
    return docx.add_heading("", level=min(max(block.level, 1), 4))


def _docx_apply_heading_run_overrides(paragraph, text: str, Pt, using_template: bool) -> None:
    if not using_template:
        return
    size = None
    stripped = text.strip()
    if stripped == "摘要":
        size = Pt(14)
    elif stripped == "关键词":
        size = Pt(12)
    if size is not None:
        for run in paragraph.runs:
            run.font.size = size


def _docx_add_list_paragraph(docx, ordered: bool, index: int):
    preferred = "List Number" if ordered else "List Bullet"
    style = _docx_first_existing_style(docx, [preferred, "List Paragraph"])
    paragraph = docx.add_paragraph(style=style) if style else docx.add_paragraph()
    if style != preferred:
        paragraph.add_run(f"{index}. " if ordered else "- ")
    return paragraph


def _docx_first_existing_style(docx, names: list[str]) -> str | None:
    for name in names:
        try:
            docx.styles[name]
        except KeyError:
            continue
        return name
    return None


def _docx_set_run_fonts(run, ascii_font: str, east_asia_font: str) -> None:
    from docx.oxml.ns import qn

    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def _docx_format_table(table, col_count: int, Inches, Pt) -> None:
    table.autofit = True
    font_size = _table_font_size(col_count, rows=len(table.rows))
    cell_width = max(0.55, min(1.25, 6.2 / max(col_count, 1)))
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(cell_width)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _docx_set_run_fonts(run, _DOCX_CJK_FONT, _DOCX_CJK_FONT)
                    run.font.size = Pt(font_size)


def _docx_add_runs(paragraph, text: str, eq_dir: Path) -> None:
    """Add inline runs to a paragraph, rendering math as small embedded images.

    No OMML/Office Math is used — all equations become PNG images via
    matplotlib, so they render correctly in Word, LibreOffice, WPS, and
    browser-based viewers without leaking raw OOXML tags.
    """
    from docx.shared import Inches

    for run in iter_inline_runs(text):
        if run.kind == "text":
            if run.content:
                paragraph.add_run(run.content)
        elif run.kind == "bold":
            r = paragraph.add_run(run.content)
            r.bold = True
        else:  # inline math
            image = render_latex(run.content, eq_dir, display=False)
            if image is not None and image.path.exists():
                try:
                    run_obj = paragraph.add_run()
                    run_obj.add_picture(
                        str(image.path),
                        width=Inches(min(image.width_pt / 72, 5.5)),
                    )
                except Exception:
                    paragraph.add_run(_math_fallback_text(run.content))
            else:
                paragraph.add_run(_math_fallback_text(run.content))


def _math_fallback_text(latex: str) -> str:
    """Minimal Unicode fallback when math rendering fails.

    Strips LaTeX commands and keeps only readable characters so the paper
    is still legible even without rendered equation images.
    """
    import re

    # Remove common LaTeX commands, keep Greek letters and math symbols
    cleaned = re.sub(r"\\[a-zA-Z]+", "", latex)
    cleaned = cleaned.replace("{", "").replace("}", "")
    cleaned = cleaned.replace("$", "").replace("\\", "")
    cleaned = cleaned.replace("_", "").replace("^", "")
    return cleaned.strip() or latex[:40]


# ---------------------------------------------------------------------------
# PDF (reportlab + Adobe STSong-Light for CJK)
# ---------------------------------------------------------------------------


def export_pdf(doc: Document, output_path: Path) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        Image as RLImage,
        ListFlowable,
        ListItem,
        Paragraph as RLParagraph,
        Preformatted,
        SimpleDocTemplate,
        Spacer,
        Table as RLTable,
        TableStyle,
    )

    pdfmetrics.registerFont(UnicodeCIDFont(_PDF_CJK_FONT))

    base = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=base["Normal"], fontName=_PDF_CJK_FONT, fontSize=10.5, leading=16)
    title_style = ParagraphStyle("DocTitle", parent=body, fontSize=20, leading=26, spaceAfter=14, alignment=TA_CENTER)
    caption_style = ParagraphStyle("Caption", parent=body, fontSize=9, textColor=colors.grey, alignment=TA_CENTER)
    code_style = ParagraphStyle(
        "Code",
        parent=base["Code"],
        fontName=_PDF_CJK_FONT,
        fontSize=8,
        leading=11,
        wordWrap="CJK",
        splitLongWords=True,
    )
    heading_styles = {
        level: ParagraphStyle(
            f"H{level}",
            parent=body,
            fontSize=max(18 - 2 * level, 11),
            leading=max(22 - 2 * level, 15),
            spaceBefore=10,
            spaceAfter=6,
        )
        for level in range(1, 5)
    }

    output_path.parent.mkdir(parents=True, exist_ok=True)
    eq_dir = _equation_dir(output_path)
    story: list = [RLParagraph(_esc(doc.title), title_style), Spacer(1, 0.3 * cm)]
    max_width = A4[0] - 4 * cm

    for block in _content_blocks(doc):
        if isinstance(block, Heading):
            level = min(max(block.level, 1), 4)
            story.append(RLParagraph(_rl_inline(block.text, eq_dir), heading_styles[level]))
        elif isinstance(block, Paragraph):
            story.append(RLParagraph(_rl_inline(block.text, eq_dir), body))
        elif isinstance(block, MathBlock):
            image = render_latex(block.latex, eq_dir, display=True)
            if image is not None and image.path.exists():
                width, height = image.width_pt, image.height_pt
                if width > max_width:
                    height *= max_width / width
                    width = max_width
                flowable = RLImage(str(image.path), width=width, height=height)
                flowable.hAlign = "CENTER"
                story.append(flowable)
            else:
                story.append(RLParagraph(_esc(block.latex), body))
        elif isinstance(block, ListBlock):
            items = [ListItem(RLParagraph(_rl_inline(item, eq_dir), body)) for item in block.items]
            story.append(
                ListFlowable(items, bulletType="1" if block.ordered else "bullet", start="1" if block.ordered else None)
            )
        elif isinstance(block, CodeBlock):
            story.append(Preformatted(_wrap_code_text(block.text), code_style))
        elif isinstance(block, ImageBlock):
            flowable = _pdf_image(block.path, max_width, RLImage, ImageReader)
            if flowable is not None:
                story.append(flowable)
                if block.caption:
                    story.append(RLParagraph(_rl_inline(block.caption, eq_dir), caption_style))
                story.append(Spacer(1, 0.2 * cm))
        elif isinstance(block, TableBlock):
            story.extend(_pdf_table_flowables(block, max_width, body, RLTable, TableStyle, colors, eq_dir))
            story.append(Spacer(1, 0.2 * cm))

    SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=doc.title,
    ).build(story)
    return output_path


def _pdf_image(path: Path, max_width: float, RLImage, ImageReader):
    if not path.exists():
        return None
    try:
        reader = ImageReader(str(path))
        iw, ih = reader.getSize()
        if iw <= 0 or ih <= 0:
            return None
        scale = min(max_width / iw, 1.0)
        return RLImage(str(path), width=iw * scale, height=ih * scale)
    except Exception:
        return None


def _rl_inline(text: str, eq_dir: Path) -> str:
    parts: list[str] = []
    for run in iter_inline_runs(text):
        if run.kind == "text":
            parts.append(_esc(run.content))
        elif run.kind == "bold":
            parts.append("<b>" + _esc(run.content) + "</b>")
        else:
            image = render_latex(run.content, eq_dir, display=False)
            if image is not None:
                parts.append(
                    f'<img src="{image.path.as_posix()}" width="{image.width_pt:.2f}" '
                    f'height="{image.height_pt:.2f}" valign="-2"/>'
                )
            else:
                parts.append(_esc(run.content))
    return "".join(parts)


def _pdf_table_flowables(block: TableBlock, max_width: float, body_style, RLTable, TableStyle, colors, eq_dir: Path):
    flowables = []
    for caption, headers, rows in _split_table_columns(block):
        flowables.append(_pdf_table(caption, headers, rows, max_width, body_style, RLTable, TableStyle, colors, eq_dir))
    return flowables


def _pdf_table(caption: str, headers: list[str], rows: list[list[str]], max_width: float, body_style, RLTable, TableStyle, colors, eq_dir: Path):
    from reportlab.platypus import Paragraph as RLParagraph
    try:
        from reportlab.platypus import LongTable
    except Exception:  # pragma: no cover - reportlab has LongTable in supported versions
        LongTable = RLTable

    cell_style = body_style.clone("Cell")
    cell_style.fontSize = _table_font_size(len(headers), rows=len(rows))
    cell_style.leading = cell_style.fontSize + 2.2
    cell_style.wordWrap = "CJK"
    cell_style.splitLongWords = True

    data = []
    repeat_rows = 1
    header_index = 0
    if caption:
        caption_style = cell_style.clone("TableCaptionCell")
        caption_style.fontSize = max(cell_style.fontSize, 8)
        caption_style.leading = caption_style.fontSize + 3
        data.append([RLParagraph(_rl_inline(caption, eq_dir), caption_style)] + [""] * (len(headers) - 1))
        repeat_rows = 2
        header_index = 1

    header_cells = [RLParagraph(_rl_inline(h, eq_dir), cell_style) for h in headers]
    data.append(header_cells)
    for row in rows:
        padded = list(row[: len(headers)])
        padded += [""] * (len(headers) - len(padded))
        data.append([RLParagraph(_rl_inline(value, eq_dir), cell_style) for value in padded])

    col_count = max(len(headers), 1)
    col_widths = _pdf_col_widths(headers, rows, max_width)
    table_cls = LongTable if len(rows) >= _LONG_TABLE_ROWS else RLTable
    try:
        table = table_cls(data, colWidths=col_widths, repeatRows=repeat_rows, splitByRow=True, splitInRow=True)
    except TypeError:
        table = table_cls(data, colWidths=col_widths, repeatRows=repeat_rows, splitByRow=True)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("BACKGROUND", (0, header_index), (-1, header_index), colors.HexColor("#3b6ea5")),
                ("TEXTCOLOR", (0, header_index), (-1, header_index), colors.white),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, header_index + 1), (-1, -1), [colors.white, colors.HexColor("#eef3f8")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    if caption:
        table.setStyle(
            TableStyle(
                [
                    ("SPAN", (0, 0), (-1, 0)),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.white),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ]
            )
        )
    return table


def _table_font_size(col_count: int, *, rows: int = 0) -> float:
    if col_count >= 7:
        return 6.8
    if col_count >= 5:
        return 7.5
    if rows >= _LONG_TABLE_ROWS:
        return 8.0
    return 8.5


def _pdf_col_widths(headers: list[str], rows: list[list[str]], max_width: float) -> list[float]:
    col_count = max(len(headers), 1)
    weights = []
    sample_rows = rows[:30]
    for idx, header in enumerate(headers):
        max_len = len(str(header))
        for row in sample_rows:
            if idx < len(row):
                max_len = max(max_len, min(len(str(row[idx])), 36))
        weights.append(max(4, min(max_len, 24)))
    total = sum(weights) or col_count
    min_width = min(max_width / col_count, 42)
    widths = [max(min_width, max_width * weight / total) for weight in weights]
    scale = max_width / sum(widths)
    return [width * scale for width in widths]


def _split_table_columns(block: TableBlock, max_columns: int = _MAX_TABLE_COLUMNS) -> list[tuple[str, list[str], list[list[str]]]]:
    headers = list(block.headers) or [""]
    rows = [(list(row) + [""] * len(headers))[: len(headers)] for row in block.rows]
    if len(headers) <= max_columns:
        return [(block.caption, headers, rows)]

    chunks: list[tuple[str, list[str], list[list[str]]]] = []
    key_header = headers[0]
    remaining = list(range(1, len(headers)))
    chunk_size = max(max_columns - 1, 1)
    total = (len(remaining) + chunk_size - 1) // chunk_size
    for chunk_no, start in enumerate(range(0, len(remaining), chunk_size), start=1):
        indices = [0] + remaining[start : start + chunk_size]
        chunk_headers = [headers[idx] for idx in indices]
        chunk_rows = [[row[idx] if idx < len(row) else "" for idx in indices] for row in rows]
        suffix = f" ({chunk_no}/{total}, key: {key_header})"
        chunks.append(((block.caption or "Table") + suffix, chunk_headers, chunk_rows))
    return chunks


def _wrap_code_text(text: str, max_chars: int = _LONG_CODE_LINE) -> str:
    import textwrap

    wrapped: list[str] = []
    for line in text.splitlines() or [""]:
        if len(line) <= max_chars:
            wrapped.append(line)
            continue
        indent = line[: len(line) - len(line.lstrip())]
        wrapped.extend(
            textwrap.wrap(
                line,
                width=max_chars,
                subsequent_indent=indent + "    ",
                break_long_words=True,
                break_on_hyphens=False,
                replace_whitespace=False,
                drop_whitespace=False,
            )
        )
    return "\n".join(wrapped)


# ---------------------------------------------------------------------------
# LaTeX (.tex, compile with xelatex for Chinese via ctex)
# ---------------------------------------------------------------------------

_LATEX_SPECIAL = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def _latex_escape(text: str) -> str:
    out: list[str] = []
    for ch in text:
        out.append(_LATEX_SPECIAL.get(ch, ch))
    return "".join(out)


def _latex_inline(text: str) -> str:
    parts: list[str] = []
    for run in iter_inline_runs(text):
        if run.kind == "text":
            parts.append(_latex_escape(run.content))
        elif run.kind == "bold":
            parts.append(r"\textbf{" + _latex_escape(run.content) + "}")
        else:
            parts.append("$" + run.content + "$")
    return "".join(parts)


def export_latex(doc: Document, output_path: Path) -> Path:
    section_cmds = {1: "section", 2: "subsection", 3: "subsubsection", 4: "paragraph"}
    lines: list[str] = [
        r"\documentclass[12pt]{article}",
        r"\usepackage{ctex}",
        r"\usepackage{graphicx}",
        r"\usepackage{booktabs}",
        r"\usepackage{geometry}",
        r"\usepackage{float}",
        r"\geometry{a4paper,margin=2.5cm}",
        r"\title{" + _latex_escape(doc.title) + "}",
        r"\date{}",
        r"\begin{document}",
        r"\maketitle",
    ]

    for block in _content_blocks(doc):
        if isinstance(block, Heading):
            cmd = section_cmds.get(min(max(block.level, 1), 4), "paragraph")
            lines.append(f"\\{cmd}*{{{_latex_inline(block.text)}}}")
        elif isinstance(block, Paragraph):
            lines.append(_latex_inline(block.text))
            lines.append("")
        elif isinstance(block, MathBlock):
            lines.append(r"\[")
            lines.append(block.latex)
            lines.append(r"\]")
            lines.append("")
        elif isinstance(block, ListBlock):
            env = "enumerate" if block.ordered else "itemize"
            lines.append(f"\\begin{{{env}}}")
            for item in block.items:
                lines.append(r"  \item " + _latex_inline(item))
            lines.append(f"\\end{{{env}}}")
        elif isinstance(block, CodeBlock):
            lines.append(r"\begin{verbatim}")
            lines.append(block.text)
            lines.append(r"\end{verbatim}")
        elif isinstance(block, ImageBlock):
            if block.path.exists():
                path_str = block.path.as_posix()
                lines.append(r"\begin{figure}[H]")
                lines.append(r"  \centering")
                lines.append(r"  \includegraphics[width=0.85\textwidth]{" + path_str + "}")
                if block.caption:
                    lines.append(r"  \caption{" + _latex_inline(block.caption) + "}")
                lines.append(r"\end{figure}")
        elif isinstance(block, TableBlock):
            lines.extend(_latex_table(block))

    lines.append(r"\end{document}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


def _latex_table(block: TableBlock) -> list[str]:
    col_count = max(len(block.headers), 1)
    col_spec = "|" + "c|" * col_count
    out = [r"\begin{table}[H]", r"  \centering"]
    if block.caption:
        out.append(r"  \caption{" + _latex_escape(block.caption) + "}")
    out.append(r"  \begin{tabular}{" + col_spec + "}")
    out.append(r"    \hline")
    out.append("    " + " & ".join(_latex_inline(h) for h in block.headers) + r" \\")
    out.append(r"    \hline")
    for row in block.rows:
        padded = list(row[:col_count]) + [""] * (col_count - len(row))
        out.append("    " + " & ".join(_latex_inline(v) for v in padded) + r" \\")
    out.append(r"    \hline")
    out.append(r"  \end{tabular}")
    out.append(r"\end{table}")
    return out


# ---------------------------------------------------------------------------
# Public dispatch
# ---------------------------------------------------------------------------

_EXPORTERS = {
    "docx": (export_docx, ".docx"),
    "pdf": (export_pdf, ".pdf"),
    "latex": (export_latex, ".tex"),
    "tex": (export_latex, ".tex"),
}

SUPPORTED_FORMATS = ("docx", "pdf", "latex")


def _esc(text: str) -> str:
    """Escape XML-sensitive characters for reportlab paragraphs."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _content_blocks(doc: Document):
    skipped_title = False
    for block in doc.blocks:
        if (
            not skipped_title
            and isinstance(block, Heading)
            and block.level == 1
            and _normalize_title(block.text) == _normalize_title(doc.title)
        ):
            skipped_title = True
            continue
        skipped_title = True
        yield block


def _normalize_title(text: str) -> str:
    return "".join(str(text).split())


def check_export_layout(doc: Document) -> list[str]:
    """Return non-blocking layout warnings before rendering export files."""

    warnings: list[str] = []
    if not DEFAULT_DOCX_TEMPLATE_PATH.exists():
        warnings.append(f"DOCX template asset is missing: {DEFAULT_DOCX_TEMPLATE_PATH}")
    elif _file_sha256(DEFAULT_DOCX_TEMPLATE_PATH) != DEFAULT_DOCX_TEMPLATE_SHA256:
        warnings.append("DOCX template asset hash differs from the registered 2025 national-contest template")
    for index, block in enumerate(_content_blocks(doc), start=1):
        if isinstance(block, CodeBlock):
            long_lines = [line_no for line_no, line in enumerate(block.text.splitlines(), start=1) if len(line) > _LONG_CODE_LINE]
            if long_lines:
                warnings.append(f"block {index}: code lines exceed {_LONG_CODE_LINE} chars and will be wrapped")
            if any(ord(ch) > 127 for ch in block.text):
                warnings.append(f"block {index}: code block contains non-ASCII text; CJK export fonts will be used")
        elif isinstance(block, TableBlock):
            if len(block.headers) > _MAX_TABLE_COLUMNS:
                warnings.append(
                    f"block {index}: table has {len(block.headers)} columns and will be split for export"
                )
            if len(block.rows) >= _LONG_TABLE_ROWS:
                warnings.append(
                    f"block {index}: table has {len(block.rows)} rows and will use multi-page table layout"
                )
    return warnings


def export_document(
    doc: Document,
    fmt: str,
    output_dir: Path,
    stem: str = "paper",
    docx_template_path: Path | None = DEFAULT_DOCX_TEMPLATE_PATH,
) -> Path:
    key = fmt.lower()
    if key not in _EXPORTERS:
        raise ValueError(f"Unsupported export format: {fmt}. Choose from {sorted(set(SUPPORTED_FORMATS))}.")
    func, suffix = _EXPORTERS[key]
    output_path = output_dir / f"{stem}{suffix}"
    if key == "docx":
        return export_docx(doc, output_path, template_path=docx_template_path)
    return func(doc, output_path)
