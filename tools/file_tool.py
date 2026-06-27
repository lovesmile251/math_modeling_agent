from __future__ import annotations

import zipfile
from io import BytesIO
from pathlib import Path
from typing import IO, Union


SUPPORTED_DATA_SUFFIXES = {".csv", ".tsv", ".xlsx", ".xls"}
SUPPORTED_PROBLEM_SUFFIXES = {".txt", ".md", ".docx", ".pdf"}
MAX_DATA_FILE_BYTES = 50 * 1024 * 1024
MAX_DOCUMENT_FILE_BYTES = 10 * 1024 * 1024
MAX_PDF_FILE_BYTES = 50 * 1024 * 1024
MAX_OFFICE_UNCOMPRESSED_BYTES = 250 * 1024 * 1024
MAX_ARCHIVE_ENTRIES = 10_000
MAX_COMPRESSION_RATIO = 200


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def read_docx_text(source: Union[Path, str, IO[bytes]]) -> str:
    """Extract plain text from a Word (.docx) document.

    `source` can be a path or any binary file-like object so the same helper
    works for both CLI files and in-memory Streamlit uploads.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("解析 Word 文件需要安装 python-docx，请运行 pip install python-docx。") from exc

    _validate_docx_source(source)
    document = Document(source)
    parts: list[str] = [para.text for para in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def read_problem_file(path: Path) -> str:
    """Read a problem statement from .txt, .md, .docx or .pdf."""
    if path.suffix.lower() == ".docx":
        return read_docx_text(path)
    if path.suffix.lower() == ".pdf":
        return read_pdf_text(path)
    return path.read_text(encoding="utf-8")


def read_pdf_text(source: Union[Path, str, IO[bytes]]) -> str:
    """Extract text from a PDF problem statement with a bounded file size."""
    if isinstance(source, (str, Path)):
        path = Path(source)
        if path.stat().st_size > MAX_PDF_FILE_BYTES:
            raise ValueError(f"PDF exceeds 50 MB limit: {path.name}")
        reader_source: Path | IO[bytes] = path
    else:
        position = source.tell() if hasattr(source, "tell") else None
        payload = source.read()
        if position is not None and hasattr(source, "seek"):
            source.seek(position)
        if len(payload) > MAX_PDF_FILE_BYTES:
            raise ValueError("PDF exceeds 50 MB limit.")
        reader_source = BytesIO(payload)
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError(
            "解析 PDF 文件需要安装 pypdf，请运行 pip install pypdf。"
        ) from exc

    reader = PdfReader(reader_source)
    parts = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("PDF contains no extractable text.")
    return text


def write_text(path: Path, content: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return path


def list_data_files(paths: list[Path]) -> list[Path]:
    files = [
        path for path in paths if path.exists() and path.suffix.lower() in SUPPORTED_DATA_SUFFIXES
    ]
    for path in files:
        validate_data_file(path)
    return files


def discover_data_files(data_dir: Path) -> list[Path]:
    if not data_dir.exists():
        return []
    files = sorted(
        path
        for path in data_dir.iterdir()
        if path.is_file() and path.suffix.lower() in SUPPORTED_DATA_SUFFIXES
    )
    for path in files:
        validate_data_file(path)
    return files


def validate_data_file(path: Path) -> None:
    size = path.stat().st_size
    if size > MAX_DATA_FILE_BYTES:
        raise ValueError(f"Data file exceeds 50 MB limit: {path.name}")
    if path.suffix.lower() == ".xlsx":
        _validate_zip_archive(path)


def _validate_docx_source(source: Union[Path, str, IO[bytes]]) -> None:
    if isinstance(source, (str, Path)):
        path = Path(source)
        if path.stat().st_size > MAX_DOCUMENT_FILE_BYTES:
            raise ValueError(f"Document exceeds 10 MB limit: {path.name}")
        _validate_zip_archive(path)
        return

    position = source.tell() if hasattr(source, "tell") else None
    payload = source.read()
    if position is not None and hasattr(source, "seek"):
        source.seek(position)
    if len(payload) > MAX_DOCUMENT_FILE_BYTES:
        raise ValueError("Document exceeds 10 MB limit.")
    _validate_zip_archive(BytesIO(payload))


def _validate_zip_archive(source: Path | BytesIO) -> None:
    try:
        with zipfile.ZipFile(source) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_ARCHIVE_ENTRIES:
                raise ValueError("Office archive contains too many entries.")
            total_size = sum(entry.file_size for entry in entries)
            if total_size > MAX_OFFICE_UNCOMPRESSED_BYTES:
                raise ValueError("Office archive expands beyond the 250 MB safety limit.")
            for entry in entries:
                compressed = max(entry.compress_size, 1)
                if entry.file_size / compressed > MAX_COMPRESSION_RATIO:
                    raise ValueError("Office archive has a suspicious compression ratio.")
    except zipfile.BadZipFile as exc:
        raise ValueError("Office document is not a valid ZIP-based file.") from exc
