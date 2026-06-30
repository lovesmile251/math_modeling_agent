from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

import pytest

from tools.exporters import (
    DEFAULT_DOCX_TEMPLATE_SHA256,
    DEFAULT_DOCX_TEMPLATE_PATH,
    check_docx_template_layout,
    check_export_layout,
    docx_template_field_mapping,
    export_document,
    export_docx,
    export_latex,
    export_pdf,
)
from tools.report_builder import (
    CodeBlock,
    Document,
    Heading,
    ImageBlock,
    ListBlock,
    MathBlock,
    Paragraph,
    TableBlock,
)


@pytest.fixture
def sample_png(tmp_path) -> Path:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    path = tmp_path / "fig.png"
    fig, ax = plt.subplots(figsize=(2, 1.5))
    ax.plot([0, 1, 2], [1, 3, 2])
    fig.savefig(path, dpi=80)
    plt.close(fig)
    return path


@pytest.fixture
def sample_document(sample_png) -> Document:
    return Document(
        title="测试论文",
        blocks=[
            Heading(level=1, text="一、引言"),
            Paragraph(text="这是一段包含特殊字符 & < > % _ 的中文正文，并包含公式 $a_{ij}=1$。"),
            MathBlock(latex=r"\rho(S)=\frac{2|E(S)|}{|S|(|S|-1)}"),
            ListBlock(items=["要点一", "要点二"], ordered=False),
            ListBlock(items=["步骤一", "步骤二"], ordered=True),
            CodeBlock(text="print('hi')"),
            ImageBlock(path=sample_png, caption="图1 示例曲线"),
            TableBlock(headers=["指标", "数值"], rows=[["均值", "1.5"], ["方差", "0.7"]], caption="表1 统计"),
        ],
    )


def test_export_docx(sample_document, tmp_path):
    out = export_docx(sample_document, tmp_path / "paper.docx")
    assert out.exists() and out.stat().st_size > 0

    from docx import Document as DocxDocument

    docx = DocxDocument(str(out))
    text = "\n".join(p.text for p in docx.paragraphs)
    assert "测试论文" in text
    assert "一、引言" in text
    assert len(docx.tables) == 1
    assert docx.tables[0].rows[0].cells[0].text == "指标"

    with ZipFile(out) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        # Math is rendered as embedded PNG images (no OMML leaks)
        assert not any("<m:oMath" in name.decode() if isinstance(name, bytes) else "<m:oMath" in name for name in package.namelist())
        assert "<m:oMath" not in document_xml  # no raw OMML tags leak
        assert any(name.startswith("word/media/") for name in package.namelist())  # images embedded


def test_national_contest_docx_template_asset_matches_registered_hash():
    assert DEFAULT_DOCX_TEMPLATE_PATH.exists()

    import hashlib

    asset_hash = hashlib.sha256(DEFAULT_DOCX_TEMPLATE_PATH.read_bytes()).hexdigest()

    assert asset_hash == DEFAULT_DOCX_TEMPLATE_SHA256


def test_docx_template_field_mapping_documents_formal_export_roles():
    mapping = docx_template_field_mapping()
    by_role = {item["role"]: item for item in mapping}

    assert {
        "title",
        "abstract_heading",
        "keywords",
        "section_heading",
        "subsection_heading",
        "body",
        "display_math",
        "table",
        "figure",
    }.issubset(by_role)
    assert by_role["title"]["source"] == "Document.title"
    assert "基于XXX模型" in by_role["title"]["template_anchor"]
    assert by_role["section_heading"]["docx_style"] == "Heading 1"


def test_export_docx_uses_national_contest_word_template(tmp_path):
    doc = Document(
        title="模板测试论文",
        blocks=[
            Heading(level=2, text="摘要"),
            Paragraph(text="这里是摘要正文。"),
            Heading(level=1, text="问题重述"),
            Paragraph(text="这里是正文。"),
        ],
    )

    out = export_docx(doc, tmp_path / "templated.docx")

    from docx import Document as DocxDocument

    template = DocxDocument(str(DEFAULT_DOCX_TEMPLATE_PATH))
    rendered = DocxDocument(str(out))
    text = "\n".join(p.text for p in rendered.paragraphs)

    assert "模板测试论文" in text
    assert "基于XXX模型的XXX问题研究" not in text
    assert rendered.sections[0].page_width == template.sections[0].page_width
    assert rendered.sections[0].page_height == template.sections[0].page_height
    assert rendered.sections[0].left_margin == template.sections[0].left_margin
    assert rendered.sections[0].right_margin == template.sections[0].right_margin
    assert rendered.paragraphs[0].alignment == template.paragraphs[0].alignment
    assert rendered.paragraphs[1].alignment == template.paragraphs[1].alignment
    assert rendered.paragraphs[1].runs[0].font.size == template.paragraphs[1].runs[0].font.size
    assert any(p.text == "问题重述" and p.style.name == "Heading 1" for p in rendered.paragraphs)

    layout = check_docx_template_layout(out)
    assert layout.passed, layout.warnings
    assert layout.template_sha256 == DEFAULT_DOCX_TEMPLATE_SHA256
    assert layout.metrics["template_placeholder_hits"] == 0
    assert layout.metrics["field_mapping_count"] >= 8


def test_export_pdf(sample_document, tmp_path):
    out = export_pdf(sample_document, tmp_path / "paper.pdf")
    assert out.exists()
    header = out.read_bytes()[:5]
    assert header == b"%PDF-"
    assert list((tmp_path / "_equations").glob("eq_*.png"))


def test_export_latex(sample_document, tmp_path):
    out = export_latex(sample_document, tmp_path / "paper.tex")
    content = out.read_text(encoding="utf-8")
    assert r"\documentclass" in content
    assert r"\usepackage{ctex}" in content
    assert r"\section*{一、引言}" in content
    # special characters must be escaped
    assert r"\&" in content and r"\%" in content and r"\_" in content
    assert r"\includegraphics" in content
    assert r"\begin{tabular}" in content


def test_export_latex_preserves_markdown_math(tmp_path):
    doc = Document(
        title="公式测试",
        blocks=[Paragraph(text=r"构建无向社交网络 \(G=(V,E)\)，模块度 \(Q=0.2576\)。")],
    )

    out = export_latex(doc, tmp_path / "paper.tex")
    content = out.read_text(encoding="utf-8")

    assert "$G=(V,E)$" in content
    assert "$Q=0.2576$" in content
    assert r"\textbackslash{}(" not in content


def test_export_docx_converts_markdown_inline_math(tmp_path):
    doc = Document(
        title="公式测试",
        blocks=[Paragraph(text=r"内部连接密度 \(\rho=\frac{2|E_C|}{|C|(|C|-1)}\)。")],
    )

    out = export_docx(doc, tmp_path / "paper.docx")
    with ZipFile(out) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")

    assert r"\rho" not in document_xml
    assert r"\(" not in document_xml
    # Math rendered as image, no OMML tags leak
    assert "<m:oMath" not in document_xml
    assert any(name.startswith("word/media/") for name in package.namelist())


def test_export_document_dispatch(sample_document, tmp_path):
    for fmt, suffix in (("docx", ".docx"), ("pdf", ".pdf"), ("latex", ".tex")):
        out = export_document(sample_document, fmt, tmp_path, stem="report")
        assert out.name == f"report{suffix}"
        assert out.exists()


def test_export_document_invalid_format(sample_document, tmp_path):
    with pytest.raises(ValueError):
        export_document(sample_document, "rtf", tmp_path)


def test_export_latex_skips_duplicate_markdown_title(tmp_path):
    doc = Document(title="同名标题", blocks=[Heading(level=1, text="同名标题"), Paragraph(text="正文。")])

    out = export_latex(doc, tmp_path / "paper.tex")
    content = out.read_text(encoding="utf-8")

    assert content.count("同名标题") == 1
    assert r"\section*{同名标题}" not in content


def test_export_handles_missing_image(tmp_path):
    doc = Document(title="无图", blocks=[ImageBlock(path=tmp_path / "nope.png", caption="缺失")])
    # None of the backends should raise when the referenced image is absent.
    assert export_docx(doc, tmp_path / "a.docx").exists()
    assert export_pdf(doc, tmp_path / "a.pdf").exists()
    assert export_latex(doc, tmp_path / "a.tex").exists()


def test_export_docx_code_block_sets_cjk_font(tmp_path):
    doc = Document(title="code", blocks=[CodeBlock(text='print("\u4e2d\u6587")')])

    out = export_docx(doc, tmp_path / "code.docx")

    with ZipFile(out) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
    assert 'w:ascii="Consolas"' in document_xml
    assert 'w:eastAsia="SimSun"' in document_xml


def test_export_docx_splits_wide_tables(tmp_path):
    headers = [f"h{i}" for i in range(10)]
    rows = [[f"r{row}c{col}" for col in range(10)] for row in range(3)]
    doc = Document(title="wide", blocks=[TableBlock(headers=headers, rows=rows, caption="wide table")])

    out = export_docx(doc, tmp_path / "wide.docx")

    from docx import Document as DocxDocument

    docx = DocxDocument(str(out))
    assert len(docx.tables) == 2
    assert docx.tables[0].rows[0].cells[0].text == "h0"
    assert docx.tables[1].rows[0].cells[0].text == "h0"


def test_export_pdf_long_wide_table_has_no_blank_pages(tmp_path):
    headers = [f"h{i}" for i in range(11)]
    rows = [[f"r{row}c{col}" for col in range(11)] for row in range(80)]
    doc = Document(title="wide pdf", blocks=[TableBlock(headers=headers, rows=rows, caption="long wide")])

    out = export_pdf(doc, tmp_path / "wide.pdf")

    from pypdf import PdfReader

    reader = PdfReader(str(out))
    assert len(reader.pages) > 1
    assert all((page.extract_text() or "").strip() for page in reader.pages)


def test_check_export_layout_warns_for_export_risks():
    headers = [f"h{i}" for i in range(9)]
    rows = [[str(row)] * 9 for row in range(24)]
    doc = Document(
        title="layout",
        blocks=[
            CodeBlock(text="x = " + "1" * 120 + "\n# \u4e2d\u6587"),
            TableBlock(headers=headers, rows=rows),
        ],
    )

    warnings = check_export_layout(doc)

    joined = "\n".join(warnings)
    assert "code lines exceed" in joined
    assert "non-ASCII text" in joined
    assert "will be split" in joined
    assert "multi-page table layout" in joined
